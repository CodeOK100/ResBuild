from flask import Flask, redirect, request, render_template, send_file, url_for, session, flash
from fpdf import FPDF
from docx import Document
from flask_sqlalchemy import SQLAlchemy
from werkzeug.security import generate_password_hash, check_password_hash
from dotenv import load_dotenv
import pymysql
import openai
import os
import re

load_dotenv()

# Flask setup
app = Flask(__name__)
app.secret_key = os.getenv("SECRET_KEY", "your_default_secret")

app.config['SQLALCHEMY_DATABASE_URI'] = 'mysql+pymysql://username:password@localhost/resume_builder'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)

# Directory setup
OUTPUT_DIR = "resumes"
os.makedirs(OUTPUT_DIR, exist_ok=True)
TEMPLATE_DIR = "templates/resume_templates"
TEMPLATE_PREVIEW_DIR = "static/template_previews"

# OpenAI API key setup
openai.api_key = os.getenv("OPENAI_API_KEY")

# Database models
class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    email = db.Column(db.String(120), unique=True, nullable=False)
    password = db.Column(db.String(200), nullable=False)
    resumes = db.relationship('Resume', backref='user', lazy=True)

class Resume(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    name = db.Column(db.String(200), nullable=False)
    content = db.Column(db.Text, nullable=False)
    format = db.Column(db.String(10), nullable=False)
    template = db.Column(db.String(200), nullable=False)

# Validators
EMAIL_REGEX = re.compile(r"[^@]+@[^@]+\.[^@]+")

def is_valid_email(email):
    return EMAIL_REGEX.match(email)

def is_strong_password(password):
    return len(password) >= 8 and any(c.isdigit() for c in password) and any(c.isupper() for c in password)

# Authentication routes
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']
        user = User.query.filter_by(email=email).first()

        if user and check_password_hash(user.password, password):
            session['user_id'] = user.id
            session['username'] = user.username
            return redirect(url_for('home'))
        else:
            flash("Invalid credentials", "error")
            return render_template('login.html')

    return render_template('login.html')

@app.route('/signup', methods=['GET', 'POST'])
def signup():
    if request.method == 'POST':
        username = request.form['name']
        email = request.form['email']
        password = request.form['password']
        confirm_password = request.form['confirm_password']

        if password != confirm_password:
            flash("Passwords do not match", "error")
            return render_template('signup.html')

        if not is_valid_email(email):
            flash("Invalid email format", "error")
            return render_template('signup.html')

        if not is_strong_password(password):
            flash("Password must be at least 8 characters long and contain an uppercase letter and a number", "error")
            return render_template('signup.html')

        existing_user = User.query.filter_by(email=email).first()
        if existing_user:
            flash("Email already registered", "error")
            return render_template('signup.html')

        hashed_password = generate_password_hash(password)
        new_user = User(username=username, email=email, password=hashed_password)
        db.session.add(new_user)
        db.session.commit()

        flash("Account created successfully. Please login.", "success")
        return redirect(url_for('login'))

    return render_template('signup.html')

@app.route('/logout')
def logout():
    session.clear()
    flash("You have been logged out.", "info")
    return redirect(url_for('login'))

@app.route('/')
def home():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    return render_template('resume.html')

@app.route('/select_template', methods=['POST'])
def select_template():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_data = request.form.to_dict()
    template_files = [
        {
            "file": file,
            "preview": f"template_previews/{file.replace('.docx', '.jpg')}"
        }
        for file in os.listdir(TEMPLATE_DIR) if file.endswith('.docx')
    ]
    return render_template('select_template.html', user_data=user_data, template_files=template_files)

def refine_inputs_with_ai(user_data):
    try:
        prompt = (
            "Refine the following resume details, correct grammatical errors, and improve the content:\n"
            f"User Input: {user_data}"
        )
        response = openai.Completion.create(
            engine="text-davinci-003",
            prompt=prompt,
            max_tokens=500
        )
        refined_data = response['choices'][0]['text'].strip()
        return eval(refined_data)
    except Exception as e:
        print(f"AI refinement error: {e}")
        return user_data

@app.route('/generate_resume', methods=['POST'])
def generate_resume():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_data = request.form.to_dict()
    template_selected = user_data.get('template')
    if not template_selected:
        return "No template selected. Please go back and select a template.", 400
    refined_data = refine_inputs_with_ai(user_data)
    return render_template('select_download_format.html', user_data=refined_data, template_selected=template_selected)

def fill_docx_template_with_ai(template_path, refined_data):
    document = Document(template_path)
    for paragraph in document.paragraphs:
        for key, value in refined_data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in refined_data.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, value)
    return document

@app.route('/download_resume', methods=['POST'])
def download_resume():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_data = request.form.to_dict()
    template_selected = user_data.pop("template")
    download_format = user_data.pop("format")
    template_path = os.path.join(TEMPLATE_DIR, template_selected)
    output_file_path = os.path.join(OUTPUT_DIR, f"My_Resume.{download_format}")
    try:
        refined_data = refine_inputs_with_ai(user_data)
        document = fill_docx_template_with_ai(template_path, refined_data)
        if download_format == "pdf":
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    pdf.multi_cell(0, 10, paragraph.text)
            pdf.output(output_file_path)
        elif download_format == "docx":
            document.save(output_file_path)
        elif download_format == "txt":
            with open(output_file_path, "w") as text_file:
                for paragraph in document.paragraphs:
                    if paragraph.text.strip():
                        text_file.write(paragraph.text + "\n")
        return send_file(output_file_path, as_attachment=True)
    except Exception as e:
        print(f"Error generating resume: {e}")
        return "An error occurred while generating the resume. Please try again.", 500

if __name__ == '__main__':
    app.run(debug=True)
